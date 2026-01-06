#!/usr/bin/env python3
"""
ABOUTME: Parses DOCX documents into text blocks using python-docx
ABOUTME: Extracts automatic numbering, splits by headings, converts tables to JSON
"""

import argparse
import hashlib
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    from numbering_resolver import NumberingResolver
    from table_extractor import TableExtractor
except ImportError:
    print("Error: Required modules not found. Ensure numbering_resolver.py and table_extractor.py are in the same directory.", file=sys.stderr)
    sys.exit(1)


def generate_content_uuid(heading: str, content: str, block_index: int) -> str:
    """
    Generate deterministic UUID from heading, content, and block position using SHA-256 hash.

    Args:
        heading: Block heading text
        content: Block content (text or JSON string for tables)
        block_index: Sequential block index in document (ensures uniqueness for duplicate content)

    Returns:
        32-character hexadecimal string (deterministic UUID)
    """
    # Combine heading and content for hashing
    if isinstance(content, list):
        # For tables, convert to JSON string for consistent hashing
        content_str = json.dumps(content, ensure_ascii=False, sort_keys=True)
    else:
        content_str = str(content)
    
    # Include block_index to ensure uniqueness for duplicate content under same heading
    combined = f"{block_index}|{heading}|{content_str}"
    return hashlib.sha256(combined.encode('utf-8')).hexdigest()[:32]


def is_heading_paragraph(para) -> tuple:
    """
    Check if paragraph is a heading by outline level or style.
    
    Returns:
        (is_heading: bool, level: int or None)
    """
    # Check outline level in paragraph XML
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        outline_lvl = pPr.find(qn('w:outlineLvl'))
        if outline_lvl is not None:
            level = int(outline_lvl.get(qn('w:val')))
            return (True, level + 1)  # Convert 0-based to 1-based
    
    # Check style name
    if para.style and para.style.name:
        style_name = para.style.name
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip())
                return (True, level)
            except ValueError:
                return (True, 1)
        if style_name in ('Title', '标题'):
            return (True, 0)
    
    return (False, None)


def extract_audit_blocks(file_path: str) -> list:
    """
    Extract text blocks from a DOCX file for auditing.
    
    Uses python-docx with custom numbering resolver to:
    1. Capture automatic numbering (list labels)
    2. Split document by headings
    3. Convert tables to JSON (2D array)
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        List of block dictionaries with heading, content, type, and metadata
    """
    doc = Document(file_path)
    resolver = NumberingResolver(file_path)
    
    blocks = []
    current_heading = "Preface/Uncategorized"
    current_heading_stack = []
    current_content = []
    
    # Iterate through document body elements (paragraphs and tables)
    body = doc._element.body
    
    for element in body:
        tag = element.tag.split('}')[-1]  # Remove namespace
        
        if tag == 'p':  # Paragraph
            # Get paragraph text
            para_text = ''
            for run in element.findall('.//w:r', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                for t in run.findall('w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if t.text:
                        para_text += t.text
            
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Get numbering label using our resolver
            label = resolver.get_label(element)
            full_text = f"{label} {para_text}".strip() if label else para_text
            
            # Check if this is a heading by outline level
            is_heading = False
            level = None
            
            pPr = element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                outline = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                if outline is not None:
                    is_heading = True
                    level = int(outline.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) + 1
            
            # Also check by style if no outline level
            if not is_heading:
                # Find matching paragraph object to check style
                para_idx = list(body).index(element)
                para_count = 0
                for p in doc.paragraphs:
                    if para_count == para_idx:
                        is_heading, level = is_heading_paragraph(p)
                        break
                    para_count += 1
            
            if is_heading and level is not None:
                # Save previous block
                if current_content:
                    content_text = "\n".join(current_content)
                    blocks.append({
                        "uuid": generate_content_uuid(current_heading, content_text, len(blocks)),
                        "heading": current_heading,
                        "content": content_text,
                        "type": "text",
                        "parent_headings": current_heading_stack[:-1] if current_heading_stack else []
                    })
                    current_content = []
                
                # Update heading stack
                current_heading_stack = current_heading_stack[:max(level - 1, 0)]
                current_heading_stack.append(full_text)
                current_heading = full_text
            else:
                current_content.append(full_text)
        
        elif tag == 'tbl':  # Table
            # Save pending text content
            if current_content:
                content_text = "\n".join(current_content)
                blocks.append({
                    "uuid": generate_content_uuid(current_heading, content_text, len(blocks)),
                    "heading": current_heading,
                    "content": content_text,
                    "type": "text",
                    "parent_headings": current_heading_stack[:-1] if current_heading_stack else []
                })
                current_content = []
            
            # Find corresponding table object
            table_idx = sum(1 for e in list(body)[:list(body).index(element)] if e.tag.endswith('tbl'))
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                table_data = TableExtractor.extract(table, numbering_resolver=resolver)
                
                table_heading = f"Table (under: {current_heading})"
                blocks.append({
                    "uuid": generate_content_uuid(table_heading, table_data, len(blocks)),
                    "heading": table_heading,
                    "content": table_data,
                    "type": "table",
                    "parent_headings": current_heading_stack[:-1] if current_heading_stack else []
                })
    
    # Save final block
    if current_content:
        content_text = "\n".join(current_content)
        blocks.append({
            "uuid": generate_content_uuid(current_heading, content_text, len(blocks)),
            "heading": current_heading,
            "content": content_text,
            "type": "text",
            "parent_headings": current_heading_stack[:-1] if current_heading_stack else []
        })
    
    return blocks


def format_table_for_display(table_data: list) -> str:
    """
    Format table data as readable text for display.

    Args:
        table_data: 2D list of cell values

    Returns:
        Formatted string representation
    """
    if not table_data:
        return "(empty table)"

    # Calculate column widths
    col_widths = []
    for col_idx in range(len(table_data[0]) if table_data else 0):
        max_width = 0
        for row in table_data:
            if col_idx < len(row):
                max_width = max(max_width, len(str(row[col_idx])))
        col_widths.append(min(max_width, 40))  # Cap at 40 chars

    lines = []
    for row in table_data:
        cells = []
        for i, cell in enumerate(row):
            width = col_widths[i] if i < len(col_widths) else 20
            cells.append(str(cell)[:width].ljust(width))
        lines.append(" | ".join(cells))

    return "\n".join(lines)


def save_blocks_jsonl(blocks: list, output_path: str):
    """
    Save blocks to JSONL format (one JSON object per line).
    Also removes existing manifest.jsonl to ensure clean resume state.

    Args:
        blocks: List of block dictionaries
        output_path: Path to output file
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        for block in blocks:
            f.write(json.dumps(block, ensure_ascii=False) + '\n')
    
    # Clean up old manifest.jsonl to prevent UUID mismatch in resume mode
    manifest_path = Path(output_path).parent / "manifest.jsonl"
    if manifest_path.exists():
        manifest_path.unlink()
        print(f"Removed existing manifest: {manifest_path}")


def save_blocks_json(blocks: list, output_path: str):
    """
    Save blocks to regular JSON format.
    Also removes existing manifest.jsonl to ensure clean resume state.

    Args:
        blocks: List of block dictionaries
        output_path: Path to output file
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump({
            "total_blocks": len(blocks),
            "blocks": blocks
        }, f, indent=2, ensure_ascii=False)
    
    # Clean up old manifest.jsonl to prevent UUID mismatch in resume mode
    manifest_path = Path(output_path).parent / "manifest.jsonl"
    if manifest_path.exists():
        manifest_path.unlink()
        print(f"Removed existing manifest: {manifest_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Parse DOCX documents into text blocks for auditing"
    )
    parser.add_argument(
        "document",
        type=str,
        help="Path to the DOCX file to parse"
    )
    parser.add_argument(
        "--output", "-o",
        type=str,
        help="Output file path (default: {document}_blocks.jsonl)"
    )
    parser.add_argument(
        "--format",
        type=str,
        choices=["jsonl", "json"],
        default="jsonl",
        help="Output format (default: jsonl)"
    )
    parser.add_argument(
        "--preview",
        action="store_true",
        help="Print preview of extracted blocks"
    )
    parser.add_argument(
        "--stats",
        action="store_true",
        help="Print statistics about the document"
    )

    args = parser.parse_args()

    # Validate input file
    doc_path = Path(args.document)
    if not doc_path.exists():
        print(f"Error: File not found: {args.document}", file=sys.stderr)
        sys.exit(1)

    if doc_path.suffix.lower() != '.docx':
        print(f"Warning: File does not have .docx extension: {args.document}", file=sys.stderr)

    # Extract blocks
    print(f"Parsing document: {args.document}")
    blocks = extract_audit_blocks(args.document)
    print(f"Extracted {len(blocks)} blocks")

    # Count by type
    text_blocks = sum(1 for b in blocks if b['type'] == 'text')
    table_blocks = sum(1 for b in blocks if b['type'] == 'table')
    print(f"  - Text blocks: {text_blocks}")
    print(f"  - Table blocks: {table_blocks}")

    # Print statistics
    if args.stats:
        print("\n--- Document Statistics ---")
        headings = set()
        total_chars = 0
        for block in blocks:
            headings.add(block['heading'])
            if block['type'] == 'text':
                total_chars += len(block['content'])
            elif block['type'] == 'table':
                total_chars += sum(len(str(cell)) for row in block['content'] for cell in row)

        print(f"Unique headings: {len(headings)}")
        print(f"Total characters: {total_chars:,}")
        print(f"Average block size: {total_chars // len(blocks) if blocks else 0:,} chars")

    # Print preview
    if args.preview:
        print("\n--- Block Preview (first 5) ---")
        for i, block in enumerate(blocks[:5]):
            print(f"\n[Block {i+1}] {block['heading']}")
            print(f"Type: {block['type']}")
            if block['type'] == 'text':
                content = block['content'][:200]
                if len(block['content']) > 200:
                    content += "..."
                print(f"Content: {content}")
            else:
                print(f"Table ({len(block['content'])} rows):")
                print(format_table_for_display(block['content'][:3]))
                if len(block['content']) > 3:
                    print("...")

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        output_path = doc_path.stem + "_blocks." + args.format

    # Save output
    if args.format == "jsonl":
        save_blocks_jsonl(blocks, output_path)
    else:
        save_blocks_json(blocks, output_path)

    print(f"\nSaved to: {output_path}")


if __name__ == "__main__":
    main()
